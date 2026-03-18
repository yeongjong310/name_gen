"""Word 문서 생성 모듈."""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

from name_gen.parser import NameEntry, PageData, lookup_happy_numbers


def _compact_paragraph(paragraph) -> None:
    """paragraph 전후 간격과 줄간격을 최소화."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(16)


def _add_bold_run(paragraph, text: str, size: Pt) -> None:
    """볼드 run을 paragraph에 추가."""
    _compact_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = size


def _format_name_text(entry: NameEntry) -> str:
    """이름 표시 텍스트 생성. 예: 김강민(金鋼? - 강철 강, 옥돌 민)"""
    surname = entry.surname
    first = entry.first_char
    second = entry.second_char

    # reading의 마지막 글자가 이름 글자 (예: "강철강" → "강")
    hangul = surname.reading[-1] if surname.reading else "?"
    hangul += first.reading[-1] if first.reading else "?"
    hangul += second.reading[-1] if second.reading else "?"

    # 뜻과 음 사이에 공백 (예: "있을유" → "있을 유")
    def _space_reading(r: str) -> str:
        if len(r) >= 2 and " " not in r:
            return r[:-1] + " " + r[-1]
        return r

    first_desc = _space_reading(first.reading)
    second_desc = _space_reading(second.reading)

    hanja = f"{surname.hanja}{first.hanja}{second.hanja}"

    return f"{hangul}({hanja} - {first_desc}, {second_desc})"


def _fill_left_cell(
    cell,
    entry: NameEntry,
    suri_numbers: list[int],
    happy_texts: dict[str, str],
    font_size: Pt,
    is_first: bool = False,
    is_last: bool = False,
    name_count: int = 3,
) -> None:
    """왼쪽 셀에 이름 + 수리오행 + 길흉 텍스트 채우기."""
    # 첫번째 이름: 빈 줄 2개 (1이름일 때는 5개), 나머지: 빈 줄 1개
    if name_count == 1:
        blank_count = 5
    elif name_count == 2:
        blank_count = 3 if is_first else 2
    else:
        blank_count = 2 if is_first else 1
    # 기존 첫 paragraph를 빈 줄로 사용
    p0 = cell.paragraphs[0]
    _compact_paragraph(p0)
    for _ in range(blank_count - 1):
        blank = cell.add_paragraph()
        _compact_paragraph(blank)

    p = cell.add_paragraph()
    _add_bold_run(p, _format_name_text(entry), font_size)

    # 수리오행
    p2 = cell.add_paragraph()
    _add_bold_run(p2, "•수리오행", font_size)
    _add_bold_run(p2, " (원격, ", font_size)
    _add_bold_run(p2, "형격", font_size)
    _add_bold_run(p2, ", ", font_size)
    _add_bold_run(p2, "이격", font_size)
    _add_bold_run(p2, ", 정격)", font_size)

    # 초년운~말년운 (숫자 + 길흉 텍스트를 한 줄로)
    jigyeok, ingyeok, oegyeok, chonggyeok = suri_numbers
    lines = [
        ("초년운", jigyeok, "jigyeok"),
        ("장년운", ingyeok, "ingyeok"),
        ("중년운", oegyeok, "oegyeok"),
        ("말년운", chonggyeok, "chonggyeok"),
    ]
    for label, num, key in lines:
        p_line = cell.add_paragraph()
        happy_text = happy_texts.get(key, "")
        if happy_text:
            text_oneline = happy_text.replace("\n\n", " ")
            _add_bold_run(p_line, f"{label} : {text_oneline}", font_size)
        else:
            _add_bold_run(p_line, f"{label} : {num}", font_size)

    # 이름 끝에 빈 줄 1개 (3이름의 마지막은 제외)
    if not (name_count == 3 and is_last):
        trailing = cell.add_paragraph()
        _compact_paragraph(trailing)


def generate_docx(page: PageData, output_path: str) -> None:
    """PageData를 기반으로 Word 문서 생성."""
    doc = Document()

    # 페이지 설정: 가로 A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Emu(10692130)   # 11.7in
    section.page_height = Emu(7560310)   # 8.3in
    section.top_margin = Emu(1224280)
    section.bottom_margin = Emu(1080135)
    section.left_margin = Emu(1440180)
    section.right_margin = Emu(1440180)

    name_count = page.name_count
    font_size = Pt(10.5) if name_count == 3 else Pt(13)
    placeholder_size = Pt(15) if name_count >= 2 else Pt(15)

    # 행 높이 결정
    if name_count == 3:
        row_height = Emu(1407795)
    elif name_count == 2:
        row_height = Emu(2399665)
    else:
        row_height = Emu(2937510)

    # 길흉 텍스트 조회
    happy_texts = {}
    if page.suri_numbers and any(page.suri_numbers):
        happy_texts = lookup_happy_numbers(page.suri_numbers)

    # 테이블 생성 (보더 없음)
    table = doc.add_table(rows=name_count, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 테이블 보더 제거
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge == "insideH":
            # 행 사이 가로 보더 유지
            el = borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "0", qn("w:color"): "auto",
            })
        else:
            el = borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none", qn("w:sz"): "0",
                qn("w:space"): "0", qn("w:color"): "auto",
            })
        borders.append(el)
    tbl_pr.append(borders)

    for idx, entry in enumerate(page.names[:name_count]):
        row = table.rows[idx]
        row.height = row_height

        # 왼쪽 셀
        left_cell = row.cells[0]
        _fill_left_cell(left_cell, entry, page.suri_numbers, happy_texts, font_size, is_first=(idx == 0), is_last=(idx == name_count - 1), name_count=name_count)

        # 오른쪽 셀: 빈 줄 1개 + placeholder
        right_cell = row.cells[1]
        p0 = right_cell.paragraphs[0]
        _compact_paragraph(p0)
        p = right_cell.add_paragraph()
        _add_bold_run(p, "ㅇ", placeholder_size)

    doc.save(output_path)
